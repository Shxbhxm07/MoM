from io import BytesIO
from typing import Optional, List, Dict
import os
import tempfile
import threading
import time
import queue
from datetime import datetime
import warnings
import json
import asyncio
import uuid
import hashlib
import uvicorn
import traceback
import re
from scipy import signal
import requests

import numpy as np
from scipy.io import wavfile

from fastapi import FastAPI, UploadFile, File, HTTPException, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import httpx
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# PDF generation
_old_md5 = hashlib.md5
def _md5(*args, **kwargs):
    kwargs.pop("usedforsecurity", None)
    return _old_md5(*args, **kwargs)
hashlib.md5 = _md5

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

warnings.filterwarnings("ignore")

# Try to import sounddevice for live recording
try:
    import sounddevice as sd
    LIVE_RECORDING_AVAILABLE = True
except ImportError:
    LIVE_RECORDING_AVAILABLE = False
    print("[WARN] sounddevice not available, live recording disabled")

# Try to import transliteration library
try:
    from indictrans import Transliterator
    trn = Transliterator(source='hin', target='eng', build_lookup=True)
    TRANSLITERATION_AVAILABLE = True
    print("[INFO] ✓ Transliteration library loaded")
except ImportError:
    TRANSLITERATION_AVAILABLE = False
    trn = None
    print("[WARN] indic-transliteration not available, transliteration disabled")

# Try to import translation library
try:
    from googletrans import Translator
    translator = Translator()
    TRANSLATION_AVAILABLE = True
    print("[INFO] ✓ Google Translator loaded")
except ImportError:
    TRANSLATION_AVAILABLE = False
    translator = None
    print("[WARN] googletrans not available, translation disabled")

# =============================================================================
# CONFIG
# =============================================================================
WHISPER_URL = "http://speech-unified:8000"
LLAMA_URL = "http://llama-summarizer:8001"
NEMO_URL = "http://nemo-diarizer:8003"

TIMEOUT = 600.0
TRANSCRIPTION_TIMEOUT = 30.0

# Audio settings
WHISPER_SAMPLE_RATE = 16000
AUDIO_CHANNELS = 1
AUDIO_BLOCKSIZE = 8192

# Legacy settings (for file upload)
SEGMENT_INTERVAL = 5.0
MIN_AUDIO_LENGTH = 0.5
SILENCE_THRESHOLD = 0.01

# Live recording settings
LIVE_CHUNK_DURATION = 2.0
LIVE_MIN_AUDIO_LENGTH = 0.8
LIVE_SILENCE_THRESHOLD = 0.005
LIVE_VAD_ENABLED = True

FILE_CACHE_SIZE_MB = 500

# Translation settings
TRANSLATION_CHUNK_SIZE = 2000
TRANSLATION_TIMEOUT = 300.0
TRANSLATION_MAX_RETRIES = 3

# =============================================================================
# LANGUAGE CONFIGURATION
# =============================================================================
SUPPORTED_LANGUAGES = {
    "all": {
        "code": None,
        "name": "Auto-Detect (All Languages)",
        "whisper_code": None
    },
    "en": {
        "code": "en",
        "name": "English",
        "whisper_code": "en"
    },
    "hi": {
        "code": "hi",
        "name": "Hindi",
        "whisper_code": "hi"
    },
    "ru": {
        "code": "ru",
        "name": "Russian",
        "whisper_code": "ru"
    },
    "he": {
        "code": "he",
        "name": "Hebrew",
        "whisper_code": "he"
    },
    "ar": {
        "code": "ar",
        "name": "Arabic",
        "whisper_code": "ar"
    }
}

DEFAULT_LANGUAGE = "none"

# =============================================================================
# DEVICE DETECTION
# =============================================================================
def detect_supported_sample_rate() -> int:
    if not LIVE_RECORDING_AVAILABLE:
        return 48000
        
    sample_rates_to_try = [48000, 44100, 32000, 24000, 22050, 16000, 8000]
    print("[INFO] Detecting supported audio sample rates...")

    try:
        device_info = sd.query_devices(kind="input")
        default_sr = int(device_info.get("default_samplerate", 48000))
        print(f"[INFO] Default device: {device_info.get('name', 'Unknown')}")
        print(f"[INFO] Default sample rate: {default_sr}")
        if default_sr not in sample_rates_to_try:
            sample_rates_to_try.insert(0, default_sr)
    except Exception as e:
        print(f"[WARN] Could not query device: {e}")

    for sr in sample_rates_to_try:
        try:
            with sd.InputStream(samplerate=sr, channels=1, dtype=np.float32, blocksize=1024):
                print(f"[INFO] ✓ Sample rate {sr} Hz is supported")
                return sr
        except Exception:
            continue

    return 48000

RECORDING_SAMPLE_RATE = 16000#detect_supported_sample_rate()

def get_audio_device_id():
    """Auto-detect EMEET microphone"""
    import sounddevice as sd
    import os
    
    # Check environment variable first
    env_device = os.getenv('AUDIO_DEVICE_ID')
    if env_device:
        try:
            return int(env_device)
        except:
            pass
    
    # Auto-detect EMEET
    for idx, device in enumerate(sd.query_devices()):
        if 'EMEET' in device['name'] and device['max_input_channels'] > 0:
            print(f"[INFO] Found EMEET at device {idx}")
            return idx
    
    return sd.default.device[0]

# Use it:
INPUT_DEVICE_ID = get_audio_device_id()

logger.info(f"{INPUT_DEVICE_ID = }")

# INPUT_DEVICE_ID = 4

# =============================================================================
# Pydantic Models
# =============================================================================
class TranscriptionResponse(BaseModel):
    formatted_text: str
    full_text: str

class LiveStatusResponse(BaseModel):
    is_recording: bool
    duration: float
    word_count: int
    total_segments: int
    latest_text: Optional[str] = None
    transcript: str
    audio_level: float

class ExportRequest(BaseModel):
    content: str
    title: Optional[str] = None

class SummarizeRequest(BaseModel):
    text: str
    max_length: int = 300

class CompleteExportRequest(BaseModel):
    formatted_transcript: str
    summary: Optional[str] = None
    raw_transcript: Optional[str] = ""
    filename: str
    speaker_count: int

class TranslateTextRequest(BaseModel):
    text: str
    source_lang: str = "hi"
    target_lang: str = "en"

class LiveStartRequest(BaseModel):
    language: str = "all"

class LanguageInfo(BaseModel):
    code: str
    name: str

class AvailableLanguagesResponse(BaseModel):
    languages: Dict[str, LanguageInfo]
    default: str

# =============================================================================
# PROGRESS TRACKING & CACHING
# =============================================================================
class ProgressTracker:
    def __init__(self):
        self.jobs: Dict[str, Dict] = {}
        self.websockets: Dict[str, List[WebSocket]] = {}
        self._lock = threading.Lock()

    def create_job(self, job_id: str):
        with self._lock:
            self.jobs[job_id] = {
                "stage": "queued",
                "progress": 0.0,
                "message": "Job queued",
                "eta_seconds": None,
                "start_time": time.time()
            }

    def update_progress(self, job_id: str, stage: str, progress: float, message: str, eta: Optional[float] = None):
        with self._lock:
            if job_id in self.jobs:
                self.jobs[job_id].update({
                    "stage": stage,
                    "progress": progress,
                    "message": message,
                    "eta_seconds": eta
                })

    def get_job_status(self, job_id: str) -> Optional[Dict]:
        with self._lock:
            return self.jobs.get(job_id)

class FileCache:
    def __init__(self, max_size_mb: int = 500):
        self.cache: Dict[str, Dict] = {}
        self.max_size = max_size_mb * 1024 * 1024
        self.current_size = 0
        self._lock = threading.Lock()

    def get_file_hash(self, file_content: bytes) -> str:
        return hashlib.sha256(file_content).hexdigest()

    def get(self, file_hash: str) -> Optional[Dict]:
        with self._lock:
            if file_hash in self.cache:
                return self.cache[file_hash]
            return None

    def set(self, file_hash: str, result: Dict, file_size: int):
        with self._lock:
            while self.current_size + file_size > self.max_size and self.cache:
                oldest_key = next(iter(self.cache))
                oldest = self.cache.pop(oldest_key)
                self.current_size -= oldest.get("file_size", 0)

            self.cache[file_hash] = {
                **result,
                "file_size": file_size,
                "cached_at": time.time()
            }
            self.current_size += file_size

    def clear(self):
        with self._lock:
            self.cache.clear()
            self.current_size = 0

progress_tracker = ProgressTracker()
file_cache = FileCache(max_size_mb=FILE_CACHE_SIZE_MB)

# =============================================================================
# SPEAKER LABEL NORMALIZATION - BULLETPROOF VERSION
# =============================================================================
def extract_speaker_number(speaker_raw: str) -> int:
    """Extract just the number from any speaker format"""
    if not speaker_raw:
        return 0
    
    # Convert to string
    s = str(speaker_raw).strip()
    
    # Find all numbers in the string
    numbers = re.findall(r'\d+', s)
    
    if numbers:
        return int(numbers[0])
    
    return 0


def make_speaker_label(number: int) -> str:
    """Create a clean speaker label"""
    return f"SPEAKER_{number}"


def fix_speaker_in_line(line: str) -> str:
    """Fix any speaker label issues in a single line"""
    if not line or 'SPEAKER' not in line.upper():
        return line
    
    # Pattern to match any variation of speaker label
    # Matches: SPEAKER_0, SPEAKER_SPEAKER_0, speaker_0, SPEAKER_SPEAKER_SPEAKER_0, etc.
    pattern = r'(?:SPEAKER_|speaker_)+(\d+)'
    
    def replace_speaker(match):
        num = match.group(1)
        return f"SPEAKER_{num}"
    
    fixed = re.sub(pattern, replace_speaker, line, flags=re.IGNORECASE)
    return fixed


def normalize_transcript_speakers(transcript: str) -> str:
    """Normalize all speaker labels in a transcript"""
    if not transcript:
        return transcript
    
    lines = transcript.split('\n')
    fixed_lines = []
    
    for line in lines:
        if line.strip():
            fixed_lines.append(fix_speaker_in_line(line))
        else:
            fixed_lines.append(line)
    
    return '\n'.join(fixed_lines)


def renumber_speakers(transcript: str) -> str:
    """Renumber speakers to 0, 1, 2... based on first appearance"""
    if not transcript:
        return transcript
    
    lines = transcript.split('\n')
    speaker_map = {}
    next_num = 0
    result_lines = []
    
    for line in lines:
        if not line.strip():
            result_lines.append(line)
            continue
        
        # Find speaker in line
        match = re.search(r'SPEAKER_(\d+)', line)
        if match:
            original_num = match.group(1)
            original_label = f"SPEAKER_{original_num}"
            
            if original_label not in speaker_map:
                speaker_map[original_label] = f"SPEAKER_{next_num}"
                next_num += 1
            
            new_label = speaker_map[original_label]
            line = line.replace(original_label, new_label)
        
        result_lines.append(line)
    
    return '\n'.join(result_lines)


# =============================================================================
# TIMESTAMP UTILITIES
# =============================================================================
def seconds_to_mmss(seconds: float) -> str:
    """Convert seconds to MM:SS format"""
    if seconds < 0:
        seconds = 0
    mins = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{mins:02d}:{secs:02d}"


def parse_timestamp_to_seconds(ts: str) -> float:
    """Parse various timestamp formats to seconds"""
    if not ts:
        return 0.0
    
    ts = ts.strip()
    parts = ts.split(':')
    
    try:
        if len(parts) == 3:  # HH:MM:SS
            h, m, s = parts
            return int(h) * 3600 + int(m) * 60 + float(s)
        elif len(parts) == 2:  # MM:SS
            m, s = parts
            return int(m) * 60 + float(s)
        else:
            return float(ts)
    except:
        return 0.0


# =============================================================================
# HALLUCINATION FILTERING
# =============================================================================
def filter_hallucinations(transcript: str) -> str:
    """Remove common Whisper hallucinations before processing"""
    hallucination_patterns = [
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*thanks for watching',
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*please give it a like',
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*subscribe to my channel',
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*if you enjoyed this video',
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*consider subscribing',
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*thank you\s*$',
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*cool\s*$',
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*gracias\s*$',
        r'\[[\d:]+\]\s*SPEAKER_\d+:\s*bye bye\s*$',
    ]
    
    lines = transcript.split('\n')
    filtered_lines = []
    
    for line in lines:
        is_hallucination = False
        line_lower = line.lower()
        
        for pattern in hallucination_patterns:
            if re.search(pattern, line_lower):
                is_hallucination = True
                print(f"[FILTERED] Hallucination: {line[:60]}...")
                break
        
        if not is_hallucination:
            filtered_lines.append(line)
    
    return '\n'.join(filtered_lines)


def clean_llama_translation_output(text: str) -> str:
    """Remove meta-text that Llama sometimes adds"""
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        lower = line.lower().strip()
        
        skip_patterns = [
            'here is the translation',
            'note:',
            'translation:',
            'i have translated',
            "i've translated",
            'the translation is',
            'translated text:',
            'output:',
        ]
        
        should_skip = any(lower.startswith(pattern) for pattern in skip_patterns)
        
        if should_skip:
            print(f"[CLEANUP] Removed meta-text: {line[:60]}...")
            continue
        
        if not line.strip():
            continue
        
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

# =============================================================================
# TRANSLATION FUNCTIONS
# =============================================================================

async def translate_chunk_with_retry(chunk: str, source_lang: str, target_lang: str, chunk_num: int, total_chunks: int, max_retries: int = 3) -> str:
    """Translate a single chunk with retry logic"""
    for attempt in range(max_retries):
        try:
            async with httpx.AsyncClient(timeout=TRANSLATION_TIMEOUT) as client:
                print(f"[TRANSLATE] Chunk {chunk_num}/{total_chunks} - Attempt {attempt + 1}/{max_retries}")
                
                response = await client.post(
                    f"{LLAMA_URL}/translate",
                    json={
                        "text": chunk,
                        "source_lang": source_lang,
                        "target_lang": target_lang,
                        "max_length": len(chunk) * 3,
                        "temperature": 0.2
                    }
                )
                
                if response.status_code == 200:
                    result = response.json()
                    translated = result.get("translated_text", "")
                    translated = clean_llama_translation_output(translated)
                    
                    if translated and len(translated.strip()) > 10:
                        return translated
                    else:
                        if attempt < max_retries - 1:
                            await asyncio.sleep(2)
                            continue
                else:
                    if attempt < max_retries - 1:
                        await asyncio.sleep(2)
                        continue
        
        except httpx.TimeoutException:
            if attempt < max_retries - 1:
                await asyncio.sleep(3)
                continue
        except Exception as e:
            if attempt < max_retries - 1:
                await asyncio.sleep(2)
                continue
    
    return chunk


async def translate_text_with_llama(text: str, source_lang: str, target_lang: str) -> str:
    """Translate text using Llama model"""
    if not text or len(text.strip()) == 0:
        return text
    
    try:
        if source_lang == target_lang:
            return text
        
        chunks = []
        if len(text) > TRANSLATION_CHUNK_SIZE:
            lines = text.split('\n')
            current_chunk = ""
            
            for line in lines:
                if len(current_chunk) + len(line) + 1 < TRANSLATION_CHUNK_SIZE:
                    current_chunk += line + "\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = line + "\n"
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        else:
            chunks = [text]
        
        translated_chunks = []
        
        for i, chunk in enumerate(chunks):
            translated = await translate_chunk_with_retry(
                chunk, source_lang, target_lang, i + 1, len(chunks),
                max_retries=TRANSLATION_MAX_RETRIES
            )
            translated_chunks.append(translated)
        
        return "\n".join(translated_chunks)
    
    except Exception as e:
        print(f"[ERROR] Translation pipeline failed: {e}")
        return text


async def translate_transcript_preserve_format(transcript: str, source_lang: str, target_lang: str) -> str:
    """Translate transcript while preserving timestamps and speaker labels"""
    try:
        transcript = filter_hallucinations(transcript)
        
        lines = transcript.strip().split('\n')
        translated_lines = []
        
        for i, line in enumerate(lines):
            if not line.strip():
                translated_lines.append("")
                continue
            
            if ']' in line and ':' in line:
                timestamp_match = line.split(']', 1)
                if len(timestamp_match) == 2:
                    timestamp = timestamp_match[0] + ']'
                    rest = timestamp_match[1].strip()
                    
                    if ':' in rest and 'SPEAKER_' in rest:
                        speaker_match = rest.split(':', 1)
                        if len(speaker_match) == 2:
                            speaker = speaker_match[0].strip()
                            text = speaker_match[1].strip()
                            
                            if len(text.strip()) < 3:
                                translated_lines.append(line)
                                continue
                            
                            try:
                                async with httpx.AsyncClient(timeout=60.0) as client:
                                    response = await client.post(
                                        f"{LLAMA_URL}/translate",
                                        json={
                                            "text": text,
                                            "source_lang": source_lang,
                                            "target_lang": target_lang,
                                            "max_length": len(text) * 3,
                                            "temperature": 0.2
                                        }
                                    )
                                    
                                    if response.status_code == 200:
                                        result = response.json()
                                        translated_text = result.get("translated_text", "").strip()
                                        translated_text = clean_llama_translation_output(translated_text)
                                        
                                        if ']' in translated_text:
                                            translated_text = translated_text.split(']', 1)[-1].strip()
                                        if 'SPEAKER_' in translated_text and ':' in translated_text:
                                            parts = translated_text.split(':', 1)
                                            if len(parts) > 1:
                                                translated_text = parts[1].strip()
                                        
                                        if '\n' in translated_text:
                                            translated_text = translated_text.split('\n')[0].strip()
                                        
                                        reconstructed = f"{timestamp} {speaker}: {translated_text}"
                                        translated_lines.append(reconstructed)
                                    else:
                                        translated_lines.append(line)
                            except Exception as e:
                                translated_lines.append(line)
                            
                            continue
            
            translated_lines.append(line)
        
        return '\n'.join(translated_lines)
    
    except Exception as e:
        print(f"[ERROR] Transcript translation failed: {e}")
        return transcript

# =============================================================================
# UTILS
# =============================================================================
async def wait_for_service(url: str, service_name: str, max_retries: int = 30, delay: int = 5) -> bool:
    """Wait for a service to become available"""
    print(f"[INFO] Waiting for {service_name} to be ready...")
    
    for attempt in range(max_retries):
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(f"{url}/health", timeout=5.0)
                if response.status_code == 200:
                    print(f"[INFO] ✓ {service_name} is ready!")
                    return True
        except Exception as e:
            if attempt < max_retries - 1:
                print(f"[INFO] {service_name} not ready yet (attempt {attempt + 1}/{max_retries}), waiting {delay}s...")
                await asyncio.sleep(delay)
            else:
                print(f"[ERROR] {service_name} failed to start: {e}")
                return False
    
    return False

def calculate_rms(audio_data: np.ndarray) -> float:
    if len(audio_data) == 0:
        return 0.0
    return float(np.sqrt(np.mean(np.square(audio_data))))

def count_words(text: str) -> int:
    return len(text.split()) if text else 0

def find_speaker_for_time(start_time: float, end_time: float, speaker_segments: list) -> str:
    """Find the speaker for a given time range"""
    mid_time = (start_time + end_time) / 2
    
    for spk_seg in speaker_segments:
        spk_start = spk_seg.get("start", 0.0)
        spk_end = spk_seg.get("end", 0.0)
        
        if spk_start <= mid_time <= spk_end:
            num = extract_speaker_number(spk_seg.get("speaker", "0"))
            return make_speaker_label(num)
    
    # Find closest
    min_dist = float('inf')
    best_num = 0
    
    for spk_seg in speaker_segments:
        spk_start = spk_seg.get("start", 0.0)
        spk_end = spk_seg.get("end", 0.0)
        
        dist = min(abs(mid_time - spk_start), abs(mid_time - spk_end))
        if dist < min_dist:
            min_dist = dist
            best_num = extract_speaker_number(spk_seg.get("speaker", "0"))
    
    return make_speaker_label(best_num)


def is_hallucination(text: str) -> bool:
    """Enhanced hallucination detection"""
    if not text or len(text.strip()) < 3:
        return False
    
    english_hallucinations = [
        'thank you', 'thanks for watching', 'subscribe', 'like and subscribe',
        'cool', 'gracias', 'bye bye', 'see you next time', 'please subscribe',
        'if you enjoyed', 'hit the bell', 'notification'
    ]
    
    hindi_hallucinations = [
        'कि', 'है', 'तो', 'और', 'यह', 'वह', 'में', 'के', 'का', 'की', 'से', 'पर',
    ]
    
    text_stripped = text.strip()
    text_lower = text.lower().strip()
    
    for phrase in english_hallucinations:
        if phrase in text_lower:
            return True
    
    if text_stripped in hindi_hallucinations:
        return True
    
    if len(text_stripped) <= 2:
        return True
    
    words = text.split()
    if len(words) >= 3:
        from collections import Counter
        word_counts = Counter(words)
        
        for word, count in word_counts.items():
            if count >= 3 and count / len(words) > 0.4:
                return True
    
    return False


def transliterate_hindi_to_english(text: str) -> str:
    """Transliterate Hindi Devanagari to Roman script"""
    if not TRANSLITERATION_AVAILABLE or not trn:
        return text
    
    try:
        if not text or len(text.strip()) == 0:
            return text
        return trn.transform(text)
    except Exception as e:
        return text


def transliterate_transcript_with_speakers(formatted_transcript: str) -> str:
    """Transliterate diarized transcript while preserving structure"""
    if not TRANSLITERATION_AVAILABLE or not trn:
        return formatted_transcript
    
    try:
        lines = formatted_transcript.split('\n')
        transliterated_lines = []
        
        for line in lines:
            if not line.strip():
                transliterated_lines.append(line)
                continue
            
            if ']' in line and ':' in line:
                timestamp_part = line.split(']', 1)[0] + ']'
                rest = line.split(']', 1)[1].strip()
                
                if ':' in rest:
                    speaker_label = rest.split(':', 1)[0]
                    text = rest.split(':', 1)[1].strip()
                    
                    transliterated_text = transliterate_hindi_to_english(text)
                    new_line = f"{timestamp_part} {speaker_label}: {transliterated_text}"
                    transliterated_lines.append(new_line)
                else:
                    transliterated_lines.append(line)
            else:
                transliterated_lines.append(transliterate_hindi_to_english(line))
        
        return '\n'.join(transliterated_lines)
    
    except Exception as e:
        return formatted_transcript


def format_diarized_transcript(transcription_result: dict, diarization_result: dict) -> str:
    """Format transcript with speaker labels"""
    formatted_lines = []
    
    transcription_segments = transcription_result.get("segments", [])
    speaker_segments = diarization_result.get("speakers", [])
    
    if not transcription_segments:
        return transcription_result.get("text", "")
    
    for trans_seg in transcription_segments:
        start_time = trans_seg.get("start", 0.0)
        end_time = trans_seg.get("end", start_time + 1.0)
        text = trans_seg.get("text", "").strip()
        
        if not text:
            continue
        
        if is_hallucination(text):
            continue
        
        speaker = find_speaker_for_time(start_time, end_time, speaker_segments)
        time_str = seconds_to_mmss(start_time)
        formatted_lines.append(f"[{time_str}] {speaker}: {text}")
    
    result = "\n".join(formatted_lines)
    result = normalize_transcript_speakers(result)
    result = renumber_speakers(result)
    
    return result


def build_pdf_buffer(content: str, title: str = "AI Pipeline Transcription") -> BytesIO:
    buffer = BytesIO()
    try:
        c = canvas.Canvas(buffer, pagesize=letter)
        c.setTitle(title)
        width, height = letter
        x, y = 50, height - 50
        
        c.setFont("Helvetica-Bold", 14)
        c.drawString(x, y, title)
        y -= 24
        
        c.setFont("Helvetica", 9)
        c.drawString(x, y, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        y -= 14
        c.drawString(x, y, f"Total Words: {count_words(content)}")
        y -= 24
        
        c.setFont("Helvetica", 10)
        for line in content.split("\n"):
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 10)
            c.drawString(x, y, line[:100])
            y -= 14
        
        c.showPage()
        c.save()
        buffer.seek(0)
        return buffer
    except Exception as e:
        buffer.seek(0)
        return buffer


# =============================================================================
# LIVE RECORDING STATE
# =============================================================================
class RecordingState:
    def __init__(self):
        self._should_stop = threading.Event()
        self._is_recording = threading.Event()
        self._audio_queue = queue.Queue(maxsize=1000)
        self._segment_queue = queue.Queue()
        self._audio_level = 0.0
        self._level_lock = threading.Lock()
        self._full_audio: List[np.ndarray] = []
        self._all_segments: List[dict] = []
        self._selected_language: str = DEFAULT_LANGUAGE
        self._language_lock = threading.Lock()
        self._start_time: Optional[float] = None  # Unix timestamp when recording started

    def reset(self):
        self._should_stop.clear()
        self._is_recording.clear()
        self._audio_level = 0.0
        while not self._audio_queue.empty():
            try:
                self._audio_queue.get_nowait()
            except:
                break
        while not self._segment_queue.empty():
            try:
                self._segment_queue.get_nowait()
            except:
                break
        self._full_audio = []
        self._all_segments = []
        self._start_time = None

    def set_language(self, language: str):
        with self._language_lock:
            if language in SUPPORTED_LANGUAGES:
                self._selected_language = language
            else:
                self._selected_language = DEFAULT_LANGUAGE

    def get_language(self) -> str:
        with self._language_lock:
            return self._selected_language

    def get_whisper_language_code(self) -> Optional[str]:
        with self._language_lock:
            lang_info = SUPPORTED_LANGUAGES.get(self._selected_language, {})
            return lang_info.get("whisper_code")

    def start(self):
        self._should_stop.clear()
        self._is_recording.set()
        self._start_time = time.time()  # Record start time as Unix timestamp

    def stop(self):
        self._should_stop.set()

    def mark_stopped(self):
        self._is_recording.clear()

    def get_elapsed_seconds(self) -> float:
        """Get seconds elapsed since recording started"""
        if self._start_time is None:
            return 0.0
        return time.time() - self._start_time

    @property
    def should_stop(self) -> bool:
        return self._should_stop.is_set()

    @property
    def is_recording(self) -> bool:
        return self._is_recording.is_set()

    def add_audio(self, audio_chunk: np.ndarray):
        try:
            self._audio_queue.put_nowait(audio_chunk)
        except queue.Full:
            try:
                self._audio_queue.get_nowait()
                self._audio_queue.put_nowait(audio_chunk)
            except:
                pass

    def add_full_audio(self, audio_chunk: np.ndarray):
        self._full_audio.append(audio_chunk.copy())

    def get_full_audio(self) -> Optional[np.ndarray]:
        if not self._full_audio:
            return None
        return np.concatenate(self._full_audio)

    def get_audio_chunks(self, timeout=0.1) -> List[np.ndarray]:
        chunks = []
        try:
            while True:
                chunk = self._audio_queue.get(timeout=timeout if not chunks else 0.001)
                chunks.append(chunk)
        except queue.Empty:
            pass
        return chunks

    def add_segment(self, segment: dict):
        self._segment_queue.put(segment)
        self._all_segments.append(segment)

    def get_all_segments(self) -> List[dict]:
        return self._all_segments.copy()

    def get_segments(self) -> List[dict]:
        segments = []
        while True:
            try:
                segments.append(self._segment_queue.get_nowait())
            except queue.Empty:
                break
        return segments

    def set_audio_level(self, level: float):
        with self._level_lock:
            self._audio_level = level

    def get_audio_level(self) -> float:
        with self._level_lock:
            return self._audio_level

rec_state = RecordingState()
recording_thread: Optional[threading.Thread] = None
transcription_thread: Optional[threading.Thread] = None
recording_start_time: Optional[datetime] = None
recording_duration_final: Optional[float] = None
live_transcript: str = ""
live_total_segments: int = 0

# =============================================================================
# LIVE RECORDING WORKERS
# =============================================================================
def recording_worker(rec_state: RecordingState, sample_rate: int, input_device=None):
    """Recording worker with device selection support"""
    if not LIVE_RECORDING_AVAILABLE:
        return
        
    def audio_callback(indata, frames, time_info, status):
        if status and "overflow" not in str(status).lower():
            print(f"[AUDIO] {status}")
        audio_data = indata[:, 0] if indata.ndim > 1 else indata.flatten()
        rms = calculate_rms(audio_data)
        rec_state.set_audio_level(rms)
        rec_state.add_audio(audio_data.copy())
        rec_state.add_full_audio(audio_data.copy())

    try:
        if input_device is not None:
            device_info = sd.query_devices(input_device, 'input')
        else:
            device_info = sd.query_devices(kind='input')
        
        device_name = device_info['name']
        
        print(f"[INFO] ═══════════════════════════════════════════════════")
        print(f"[INFO] Recording started: {device_name}")
        print(f"[INFO] Sample Rate: {sample_rate} Hz")
        print(f"[INFO] ═══════════════════════════════════════════════════")
        
        with sd.InputStream(
            channels=AUDIO_CHANNELS,
            samplerate=sample_rate,
            callback=audio_callback,
            dtype=np.float32,
            blocksize=AUDIO_BLOCKSIZE,
            latency="high",
            device=input_device,
        ):
            while not rec_state.should_stop:
                time.sleep(0.05)
            print("[INFO] Live recording stopped")
    except Exception as e:
        print(f"[ERROR] Recording error: {e}")
        traceback.print_exc()
    finally:
        rec_state.mark_stopped()


def transcription_worker(rec_state: RecordingState, sample_rate: int, segment_interval: float):
    """Transcription worker that creates segments with RELATIVE timestamps"""
    audio_buffer = []
    buffer_samples = 0
    last_transcription_time = time.time()
    segment_counter = 0

    print(f"[INFO] Transcription worker started (interval: {segment_interval}s)")

    while True:
        chunks = rec_state.get_audio_chunks(timeout=0.05)
        for chunk in chunks:
            audio_buffer.append(chunk)
            buffer_samples += len(chunk)

        current_time = time.time()
        time_elapsed = current_time - last_transcription_time

        should_transcribe = (
            (time_elapsed >= segment_interval and buffer_samples > 0) or
            (not rec_state.is_recording and buffer_samples > 0)
        )

        if should_transcribe and audio_buffer:
            audio_segment = np.concatenate(audio_buffer)
            audio_buffer = []
            buffer_samples = 0
            last_transcription_time = current_time

            min_samples = int(sample_rate * LIVE_MIN_AUDIO_LENGTH)
            if len(audio_segment) < min_samples:
                continue

            rms = calculate_rms(audio_segment)
            if rms < LIVE_SILENCE_THRESHOLD:
                continue

            segment_counter += 1
            
            # Get elapsed seconds since recording started
            elapsed = rec_state.get_elapsed_seconds()
            
            threading.Thread(
                target=process_live_transcription_segment,
                args=(audio_segment, segment_counter, rec_state, elapsed),
                daemon=True
            ).start()

        if not rec_state.is_recording and len(audio_buffer) == 0:
            break

        time.sleep(0.02)

    print("[INFO] Transcription worker stopped")


def fallback_single_speaker_diarization(existing_segments: list, generate_summary: bool):
    """Fallback when NeMo diarization fails"""
    logger.info("[FALLBACK] Using single-speaker mode")
    
    formatted_transcript = ""
    
    for seg in existing_segments:
        if is_hallucination(seg['text']):
            continue
        
        ts = seg.get('timestamp', '00:00')
        formatted_transcript += f"[{ts}] SPEAKER_0: {seg['text']}\n"
    
    summary_result = None
    
    if generate_summary:
        try:
            summary_response = requests.post(
                f"{LLAMA_URL}/summarize",
                json={"text": formatted_transcript, "max_length": 4000, "temperature": 0.01},
                timeout=300
            )
            
            if summary_response.status_code == 200:
                summary_result = summary_response.json()
        except Exception as e:
            logger.error(f"[FALLBACK] Summary failed: {e}")
    
    return {
        "formatted_transcript": formatted_transcript,
        "num_speakers": 1,
        "summary": summary_result.get('analysis') if summary_result else None,
        "segments_count": len(existing_segments),
        "fallback_mode": True
    }


def process_live_transcription_segment(audio_segment: np.ndarray, segment_id: int, rec_state: RecordingState, elapsed_seconds: float):
    """
    Transcription with RELATIVE timestamps (MM:SS from start of recording)
    """
    try:
        selected_language = rec_state.get_language()
        whisper_language_code = rec_state.get_whisper_language_code()
        
        # Resample to 16kHz for Whisper
        if RECORDING_SAMPLE_RATE != WHISPER_SAMPLE_RATE:
            num_samples = int(len(audio_segment) * WHISPER_SAMPLE_RATE / RECORDING_SAMPLE_RATE)
            audio_resampled = signal.resample(audio_segment, num_samples)
        else:
            audio_resampled = audio_segment
        
        rms_original = calculate_rms(audio_resampled)
        
        if rms_original < LIVE_SILENCE_THRESHOLD:
            return
        
        max_val = np.max(np.abs(audio_resampled))
        if max_val > 0:
            audio_normalized = audio_resampled / max_val
        else:
            return
        
        audio_int16 = (audio_normalized * 32767).astype(np.int16)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            wavfile.write(tmp.name, WHISPER_SAMPLE_RATE, audio_int16)
            tmp_path = tmp.name

        base_whisper_params = {
            'beam_size': '5',
            'best_of': '5',
            'temperature': '0.0',
            'condition_on_previous_text': 'true',
            'word_timestamps': 'true',
            'task': 'transcribe',
            'compression_ratio_threshold': '2.0',
            'logprob_threshold': '-0.8',
            'no_speech_threshold': '0.7',
        }
        
        final_text = ""
        final_language = selected_language
        
        if whisper_language_code is not None:
            with open(tmp_path, 'rb') as f:
                files = {'file': ('audio.wav', f, 'audio/wav')}
                transcribe_data = base_whisper_params.copy()
                transcribe_data['language'] = whisper_language_code
                
                response = requests.post(
                    f"{WHISPER_URL}/transcribe",
                    files=files,
                    data=transcribe_data,
                    timeout=TRANSCRIPTION_TIMEOUT
                )
                
                if response.status_code == 200:
                    result = response.json()
                    final_text = result.get('text', '').strip()
                    final_language = whisper_language_code
                else:
                    os.unlink(tmp_path)
                    return
        else:
            with open(tmp_path, 'rb') as f:
                files = {'file': ('audio.wav', f, 'audio/wav')}
                
                detect_response = requests.post(
                    f"{WHISPER_URL}/transcribe",
                    files=files,
                    data=base_whisper_params,
                    timeout=TRANSCRIPTION_TIMEOUT
                )
                
                if detect_response.status_code != 200:
                    os.unlink(tmp_path)
                    return
                
                detect_result = detect_response.json()
                detected_language = detect_result.get('language', 'en')
                detected_text = detect_result.get('text', '').strip()
            
            final_text = detected_text
            final_language = detected_language
        
        # Validation
        if not final_text or len(final_text) < 3:
            os.unlink(tmp_path)
            return
        
        if is_hallucination(final_text):
            os.unlink(tmp_path)
            return
        
        if not any(c.isalnum() for c in final_text):
            os.unlink(tmp_path)
            return
        
        # =================================================================
        # CREATE RELATIVE TIMESTAMP (MM:SS format)
        # =================================================================
        timestamp = seconds_to_mmss(elapsed_seconds)
        
        rec_state.add_segment({
            "timestamp": timestamp,
            "text": final_text,
            "language": final_language,
            "segment_id": segment_id,
            "elapsed_seconds": elapsed_seconds  # Store raw seconds for diarization
        })
        
        print(f"[LIVE] [{timestamp}] {final_text[:60]}...")
        
        os.unlink(tmp_path)
        
    except requests.Timeout:
        print(f"[TIMEOUT] Seg {segment_id}")
    except Exception as e:
        print(f"[ERROR] Seg {segment_id}: {e}")


# =============================================================================
# FASTAPI APP
# =============================================================================
app = FastAPI(
    title="AI Pipeline Gateway",
    version="2.2.0",
    description="Unified API for Whisper + NeMo + Llama"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =============================================================================
# STARTUP EVENT
# =============================================================================
@app.on_event("startup")
async def startup_event():
    print("[INFO] Starting AI Gateway...")
    
    services = [
        (WHISPER_URL, "Whisper"),
        (NEMO_URL, "NeMo"),
        (LLAMA_URL, "Llama"),
    ]
    
    for url, name in services:
        await wait_for_service(url, name)
    
    print("[INFO] ✓ All services ready!")

# =============================================================================
# ROUTES
# =============================================================================
@app.get("/")
async def root():
    return {"service": "AI Pipeline Gateway", "version": "2.2.0", "status": "ready"}

@app.get("/health")
async def health():
    return {"gateway": "healthy", "live_recording_available": LIVE_RECORDING_AVAILABLE}

@app.post("/transcribe")
async def transcribe(file: UploadFile = File(...)):
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            files = {"file": (file.filename, await file.read(), file.content_type)}
            response = await client.post(f"{WHISPER_URL}/transcribe", files=files)
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail=response.text)
            return response.json()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/diarize")
async def diarize(file: UploadFile = File(...)):
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            files = {"file": (file.filename, await file.read(), file.content_type)}
            response = await client.post(f"{NEMO_URL}/diarize", files=files)
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail=response.text)
            return response.json()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/summarize")
async def summarize(request: SummarizeRequest):
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            response = await client.post(
                f"{LLAMA_URL}/summarize",
                json={"text": request.text, "max_length": request.max_length}
            )
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail=response.text)
            return response.json()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/translate-text")
async def translate_text_endpoint(request: TranslateTextRequest):
    try:
        if not request.text or len(request.text.strip()) == 0:
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        
        cleaned_text = filter_hallucinations(request.text)
        
        if not cleaned_text or len(cleaned_text.strip()) == 0:
            raise HTTPException(status_code=400, detail="No valid content after filtering")
        
        has_speaker_labels = "[" in cleaned_text and "SPEAKER_" in cleaned_text
        
        if has_speaker_labels:
            translated = await translate_transcript_preserve_format(
                cleaned_text, request.source_lang, request.target_lang
            )
        else:
            translated = await translate_text_with_llama(
                cleaned_text, request.source_lang, request.target_lang
            )
        
        return {
            "translated_text": translated,
            "source_lang": request.source_lang,
            "target_lang": request.target_lang
        }
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")


@app.post("/full-pipeline")
async def full_pipeline(file: UploadFile = File(...), summarize_output: bool = True, transliterate: bool = True):
    file_content = await file.read()
    file_hash = file_cache.get_file_hash(file_content)
    
    cached = file_cache.get(file_hash)
    if cached:
        return {"filename": file.filename, "from_cache": True, **cached}
    
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            transcribe_task = client.post(
                f"{WHISPER_URL}/transcribe",
                files={"file": (file.filename, file_content, file.content_type)},
            )
            
            diarize_task = client.post(
                f"{NEMO_URL}/diarize",
                files={"file": (file.filename, file_content, file.content_type)},
            )
            
            results = await asyncio.gather(transcribe_task, diarize_task, return_exceptions=True)
            transcribe_response, diarize_response = results
            
            if isinstance(transcribe_response, Exception):
                raise HTTPException(500, f"Transcription failed: {str(transcribe_response)}")
            
            if isinstance(diarize_response, Exception):
                raise HTTPException(500, f"Diarization failed: {str(diarize_response)}")
            
            if transcribe_response.status_code != 200:
                raise HTTPException(transcribe_response.status_code, transcribe_response.text)
            
            if diarize_response.status_code != 200:
                raise HTTPException(diarize_response.status_code, diarize_response.text)
            
            transcription = transcribe_response.json()
            diarization = diarize_response.json()
            
            formatted_transcript = format_diarized_transcript(transcription, diarization)
            
            speakers = diarization.get("speakers", [])
            unique_speakers = set(extract_speaker_number(spk.get("speaker", "0")) for spk in speakers)
            num_speakers = len(unique_speakers) if unique_speakers else 1
            
            summary = None
            if summarize_output and transcription.get("text"):
                try:
                    summarize_response = await client.post(
                        f"{LLAMA_URL}/summarize",
                        json={"text": transcription["text"], "max_length": 300},
                    )
                    if summarize_response.status_code == 200:
                        summary = summarize_response.json()
                except Exception as e:
                    summary = {"summary": "Summarization failed", "analysis": str(e)}
            
            raw_transcript = transcription.get("text", "")
            
            formatted_transcript_transliterated = None
            raw_transcript_transliterated = None
            
            if transliterate and TRANSLITERATION_AVAILABLE:
                formatted_transcript_transliterated = transliterate_transcript_with_speakers(formatted_transcript)
                raw_transcript_transliterated = transliterate_hindi_to_english(raw_transcript)
            
            result = {
                "summary": summary,
                "formatted_transcript": formatted_transcript,
                "formatted_transcript_transliterated": formatted_transcript_transliterated,
                "raw_transcript": raw_transcript,
                "raw_transcript_transliterated": raw_transcript_transliterated,
                "transcription": transcription,
                "diarization": diarization,
                "num_speakers": num_speakers
            }
            
            file_cache.set(file_hash, result, len(file_content))
            
            return {"filename": file.filename, "from_cache": False, **result}
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Pipeline failed: {str(e)}")

# =============================================================================
# LIVE RECORDING ENDPOINTS
# =============================================================================

@app.get("/live/languages")
async def get_available_languages():
    languages = {}
    for code, info in SUPPORTED_LANGUAGES.items():
        languages[code] = {"code": code, "name": info["name"]}
    
    return {
        "languages": languages,
        "default": DEFAULT_LANGUAGE,
        "supported_codes": list(SUPPORTED_LANGUAGES.keys())
    }


@app.post("/live/start")
async def live_start(language: str = "all"):
    global recording_thread, transcription_thread, recording_start_time
    global recording_duration_final, live_transcript, live_total_segments
    
    if not LIVE_RECORDING_AVAILABLE:
        raise HTTPException(503, "Live recording not available")
    
    if rec_state.is_recording:
        raise HTTPException(400, "Already recording")
    
    if language not in SUPPORTED_LANGUAGES:
        raise HTTPException(400, f"Unsupported language '{language}'")

    rec_state.reset()
    rec_state.set_language(language)
    rec_state.start()

    recording_thread = threading.Thread(
        target=recording_worker,
        args=(rec_state, RECORDING_SAMPLE_RATE, INPUT_DEVICE_ID),
        daemon=True
    )
    transcription_thread = threading.Thread(
        target=transcription_worker,
        args=(rec_state, RECORDING_SAMPLE_RATE, LIVE_CHUNK_DURATION),
        daemon=True
    )
    
    recording_thread.start()
    transcription_thread.start()

    recording_start_time = datetime.now()
    recording_duration_final = None
    live_transcript = ""
    live_total_segments = 0

    return {
        "status": "recording_started",
        "language": language,
        "language_name": SUPPORTED_LANGUAGES[language]["name"]
    }


@app.post("/live/stop")
async def live_stop():
    global recording_duration_final, recording_start_time

    if not rec_state.is_recording:
        raise HTTPException(400, "Not recording")

    rec_state.stop()

    if recording_start_time:
        recording_duration_final = (datetime.now() - recording_start_time).total_seconds()

    if transcription_thread and transcription_thread.is_alive():
        transcription_thread.join(timeout=10.0)

    return {
        "status": "stopped",
        "duration": recording_duration_final,
        "total_segments": live_total_segments
    }


@app.post("/live/change-language")
async def live_change_language(language: str):
    if language not in SUPPORTED_LANGUAGES:
        raise HTTPException(400, f"Unsupported language '{language}'")
    
    old_language = rec_state.get_language()
    rec_state.set_language(language)
    
    return {
        "status": "language_changed",
        "old_language": old_language,
        "new_language": language
    }


@app.get("/live/current-language")
async def get_current_language():
    current = rec_state.get_language()
    return {
        "language": current,
        "language_name": SUPPORTED_LANGUAGES.get(current, {}).get("name", "Unknown"),
        "is_auto_detect": current == "all"
    }


@app.get("/live/status", response_model=LiveStatusResponse)
async def live_status():
    global live_transcript, live_total_segments

    new_segments = rec_state.get_segments()
    latest_text = None
    
    for seg in new_segments:
        live_transcript += f"[{seg['timestamp']}] {seg['text']}\n"
        live_total_segments += 1
        latest_text = seg['text']

    elapsed = (
        (datetime.now() - recording_start_time).total_seconds()
        if recording_start_time and rec_state.is_recording
        else (recording_duration_final or 0.0)
    )

    return LiveStatusResponse(
        is_recording=rec_state.is_recording,
        duration=elapsed,
        word_count=count_words(live_transcript),
        total_segments=live_total_segments,
        latest_text=latest_text,
        transcript=live_transcript,
        audio_level=rec_state.get_audio_level()
    )


@app.get("/live/status-extended")
async def live_status_extended():
    global live_transcript, live_total_segments

    new_segments = rec_state.get_segments()
    latest_text = None
    
    for seg in new_segments:
        live_transcript += f"[{seg['timestamp']}] {seg['text']}\n"
        live_total_segments += 1
        latest_text = seg['text']

    elapsed = (
        (datetime.now() - recording_start_time).total_seconds()
        if recording_start_time and rec_state.is_recording
        else (recording_duration_final or 0.0)
    )

    current_language = rec_state.get_language()

    return {
        "is_recording": rec_state.is_recording,
        "duration": elapsed,
        "word_count": count_words(live_transcript),
        "total_segments": live_total_segments,
        "latest_text": latest_text,
        "transcript": live_transcript,
        "audio_level": rec_state.get_audio_level(),
        "language": {
            "code": current_language,
            "name": SUPPORTED_LANGUAGES.get(current_language, {}).get("name", "Unknown"),
            "is_auto_detect": current_language == "all"
        }
    }


@app.post("/live/diarize")
async def live_diarize(generate_summary: bool = True, transliterate: bool = True):
    """Diarize live recording - FIXED VERSION"""
    global live_transcript
    
    if rec_state.is_recording:
        raise HTTPException(400, "Cannot diarize while recording is active")
    
    full_audio = rec_state.get_full_audio()
    if full_audio is None or len(full_audio) == 0:
        raise HTTPException(400, "No audio recorded")
    
    logger.info(f"[DIARIZE] Starting diarization on {len(full_audio)} samples")
    
    # Preprocess audio
    audio_data = full_audio.astype(np.float32)
    max_val = np.max(np.abs(audio_data))
    if max_val > 0:
        audio_data = audio_data / max_val * 0.7
    
    audio_data = np.clip(audio_data, -1.0, 1.0)
    audio_int16 = (audio_data * 32767).astype(np.int16)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
        wavfile.write(tmp.name, RECORDING_SAMPLE_RATE, audio_int16)
        audio_path = tmp.name
    
    try:
        existing_segments = rec_state.get_all_segments()
        
        if not existing_segments:
            raise HTTPException(400, "No transcript segments available")
        
        logger.info(f"[DIARIZE] Processing {len(existing_segments)} transcript segments")
        
        # Call NeMo diarizer
        with open(audio_path, 'rb') as f:
            files = {'file': ('audio.wav', f, 'audio/wav')}
            data = {'num_speakers': '', 'vad_onset': '0.5', 'vad_offset': '0.5'}
            
            diarization_response = requests.post(
                f"{NEMO_URL}/diarize",
                files=files,
                data=data,
                timeout=300
            )
        
        if diarization_response.status_code != 200:
            logger.error(f"[DIARIZE] NeMo failed: {diarization_response.status_code}")
            return fallback_single_speaker_diarization(existing_segments, generate_summary)
        
        diarization_result = diarization_response.json()
        
        # Find speaker segments in response
        speaker_segments = None
        for key in ['segments', 'speakers', 'diarization']:
            if key in diarization_result:
                data = diarization_result[key]
                if isinstance(data, list) and data:
                    speaker_segments = data
                    break
                elif isinstance(data, dict):
                    for sub_key in ['segments', 'speakers']:
                        if sub_key in data and data[sub_key]:
                            speaker_segments = data[sub_key]
                            break
        
        if not speaker_segments:
            logger.error(f"[DIARIZE] No speaker segments found")
            return fallback_single_speaker_diarization(existing_segments, generate_summary)
        
        logger.info(f"[DIARIZE] ✓ Found {len(speaker_segments)} speaker segments")
        
        # Build diarized transcript
        formatted_lines = []
        
        for seg in existing_segments:
            if is_hallucination(seg['text']):
                continue
            
            # Get segment time (use elapsed_seconds if available)
            seg_time = seg.get('elapsed_seconds', 0.0)
            if seg_time == 0.0 and 'timestamp' in seg:
                seg_time = parse_timestamp_to_seconds(seg['timestamp'])
            
            # Find speaker for this time
            speaker_num = 0
            for spk_seg in speaker_segments:
                start = spk_seg.get('start', 0.0)
                end = spk_seg.get('end', 0.0)
                
                if start <= seg_time <= end:
                    speaker_num = extract_speaker_number(spk_seg.get('speaker', '0'))
                    break
            
            # Build line with CLEAN speaker label
            timestamp = seg.get('timestamp', seconds_to_mmss(seg_time))
            speaker_label = make_speaker_label(speaker_num)
            
            formatted_lines.append(f"[{timestamp}] {speaker_label}: {seg['text']}")
        
        # Join and do final cleanup
        formatted_transcript = '\n'.join(formatted_lines)
        
        # Apply final normalization (renumber to 0, 1, 2...)
        formatted_transcript = renumber_speakers(formatted_transcript)
        
        logger.info(f"[DIARIZE] ✓ Built transcript with {len(formatted_lines)} lines")
        
        # Update global transcript
        live_transcript = formatted_transcript
        
        # Generate summary
        summary_result = None
        if generate_summary:
            try:
                logger.info("[DIARIZE] Generating MoM...")
                summary_response = requests.post(
                    f"{LLAMA_URL}/summarize",
                    json={"text": formatted_transcript, "max_length": 4000, "temperature": 0.01},
                    timeout=300
                )
                
                if summary_response.status_code == 200:
                    summary_result = summary_response.json()
                    logger.info("[DIARIZE] ✓ MoM generated")
            except Exception as e:
                logger.error(f"[DIARIZE] Summary failed: {e}")
        
        os.unlink(audio_path)
        
        # Count unique speakers
        unique_speakers = len(set(re.findall(r'SPEAKER_(\d+)', formatted_transcript)))
        
        logger.info(f"[DIARIZE] ✓ Complete! {unique_speakers} speakers")
        
        return {
            "formatted_transcript": formatted_transcript,
            "num_speakers": unique_speakers,
            "summary": summary_result.get('analysis') if summary_result else None,
            "segments_count": len(formatted_lines),
            "language": rec_state.get_language()
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[DIARIZE] Failed: {e}")
        traceback.print_exc()
        
        if os.path.exists(audio_path):
            os.unlink(audio_path)
        
        raise HTTPException(500, f"Diarization failed: {str(e)}")


@app.post("/live/summarize")
async def live_summarize():
    global live_transcript
    
    if rec_state.is_recording:
        raise HTTPException(400, "Stop recording first")
    
    if not live_transcript:
        raise HTTPException(400, "No transcript available")
    
    try:
        lines = live_transcript.strip().split('\n')
        text_only = [
            line.split(':', 1)[1].strip() if ':' in line and ']' in line else line
            for line in lines
        ]
        combined_text = ' '.join(text_only)
        
        async with httpx.AsyncClient(timeout=300.0) as client:
            response = await client.post(
                f"{LLAMA_URL}/summarize",
                json={"text": combined_text, "max_length": 600}
            )
            
            if response.status_code == 200:
                summary_data = response.json()
                return {
                    "status": "completed",
                    "summary": summary_data.get("analysis", summary_data.get("summary", ""))
                }
            else:
                raise HTTPException(response.status_code, response.text)
    
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/live/reset")
async def live_reset():
    global live_transcript, live_total_segments, recording_start_time, recording_duration_final
    
    if rec_state.is_recording:
        raise HTTPException(400, "Stop recording first")
    
    rec_state.reset()
    live_transcript = ""
    live_total_segments = 0
    recording_start_time = None
    recording_duration_final = None
    
    return {"status": "reset"}


@app.post("/cache/clear")
async def clear_cache():
    file_cache.clear()
    return {"status": "cleared"}

@app.get("/cache/stats")
async def cache_stats():
    return {
        "size_mb": file_cache.current_size / 1024 / 1024,
        "files": len(file_cache.cache)
    }

@app.post("/export/pdf")
async def export_pdf(req: ExportRequest):
    if not req.content:
        raise HTTPException(400, "No content")
    buffer = build_pdf_buffer(req.content, req.title or "Transcript")
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": 'attachment; filename="transcript.pdf"'})


@app.post("/export/complete-txt")
async def export_complete_txt(req: CompleteExportRequest):
    content_parts = [
        "=" * 80,
        "MEETING REPORT",
        "=" * 80,
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"File: {req.filename}",
        f"Speakers: {req.speaker_count}",
        "=" * 80,
        "",
        "TRANSCRIPT",
        "-" * 80,
        req.formatted_transcript,
        "",
        "=" * 80,
        "MINUTES OF MEETING",
        "-" * 80,
        req.summary if req.summary else "No MoM available",
        "=" * 80
    ]
    
    buffer = BytesIO("\n".join(content_parts).encode('utf-8'))
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{req.filename}_report.txt"'}
    )


@app.post("/export/complete-docx")
async def export_complete_docx(req: CompleteExportRequest):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        title = doc.add_heading('MINUTES OF MEETING', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}")
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(100, 116, 139)
        
        if req.summary and req.summary.strip():
            for line in req.summary.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)
        
        doc.add_page_break()
        
        doc.add_heading('ANNEXURE', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for line in req.formatted_transcript.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{req.filename}_report.docx"'}
        )
        
    except Exception as e:
        raise HTTPException(500, f"DOCX generation failed: {str(e)}")


@app.post("/export/complete-pdf")
async def export_complete_pdf(req: CompleteExportRequest):
    buffer = BytesIO()
    
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.75*inch, leftMargin=0.75*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, textColor=colors.HexColor('#1e40af'), alignment=TA_CENTER, fontName='Helvetica-Bold')
        body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9, leading=13, textColor=colors.HexColor('#334155'))
        transcript_style = ParagraphStyle('Transcript', parent=styles['Normal'], fontSize=8, leading=11, textColor=colors.HexColor('#1e293b'), fontName='Courier')
        
        story.append(Paragraph("MINUTES OF MEETING", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        if req.summary and req.summary.strip():
            for line in req.summary.split('\n'):
                if line.strip():
                    safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe_line, body_style))
                    story.append(Spacer(1, 4))
        
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("ANNEXURE", ParagraphStyle('AnnexureTitle', parent=styles['Heading1'], fontSize=13, alignment=TA_CENTER)))
        story.append(Spacer(1, 0.1*inch))
        
        for line in req.formatted_transcript.split('\n'):
            if line.strip():
                safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_line, transcript_style))
        
        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{req.filename}_report.pdf"'}
        )
        
    except Exception as e:
        raise HTTPException(500, f"PDF generation failed: {str(e)}")


# =============================================================================
# SERVER STARTUP
# =============================================================================
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8002)